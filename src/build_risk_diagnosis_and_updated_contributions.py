from __future__ import annotations

import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import shap
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE = Path(r"D:\VC code\MDPI2")
TABLES = BASE / "outputs" / "tables"
FIGURES = BASE / "outputs" / "figures"
OUT = BASE / "outputs" / "WasteManagement_updated_experiment_contribution_framework.docx"

sys.path.insert(0, str(BASE / "src"))
import build_empirical_pipeline as bep  # noqa: E402


FEATURE_GROUPS = {
    "recycling trajectory deficit": ["recycling_rate_lag1", "recycling_rate_lag2", "recycling_rate_trend_3y"],
    "landfill lock-in": ["landfill_rate_proxy_lag1", "landfill_rate_proxy_lag2", "landfill_rate_trend_3y"],
    "incineration/pathway pressure": [
        "incineration_rate_proxy_lag1",
        "incineration_rate_proxy_lag2",
        "incineration_rate_trend_3y",
    ],
    "circular-capacity/resource constraint": [
        "circular_material_use_rate_lag1",
        "circular_material_use_rate_lag2",
        "resource_productivity_lag1",
        "resource_productivity_lag2",
    ],
    "waste-generation and socioeconomic context": [
        "waste_generated_kg_pc_lag1",
        "waste_generated_kg_pc_lag2",
        "environmental_tax_gdp_lag1",
        "environmental_tax_gdp_lag2",
        "gdp_pc_eur_lag1",
        "gdp_pc_eur_lag2",
        "household_consumption_pc_eur_lag1",
        "household_consumption_pc_eur_lag2",
        "population_density_lag1",
        "population_density_lag2",
    ],
}


def fmt(value, digits=3):
    if value is None or pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def severity(row) -> str:
    if row["risk_probability"] < 0.5 and row["predicted_gap"] >= 0:
        return "stable achiever"
    if row["risk_probability"] < 0.5:
        return "watch list below threshold"
    if row["predicted_gap"] <= -15:
        return "high-gap high-risk"
    if row["risk_probability"] >= 0.75:
        return "high-risk near target"
    return "moderate-risk transition"


def dominant_group(shap_row: pd.Series) -> tuple[str, float]:
    scores: dict[str, float] = {}
    for group, features in FEATURE_GROUPS.items():
        # For a recycling-rate forecast, negative SHAP values reduce predicted
        # recycling and therefore increase non-compliance risk.
        vals = [float(shap_row.get(f"shap_{f}", 0.0)) for f in features if f"shap_{f}" in shap_row.index]
        scores[group] = sum(abs(v) for v in vals if v < 0)
    if not scores or max(scores.values()) <= 0:
        return "no dominant negative driver", 0.0
    group = max(scores, key=scores.get)
    return group, scores[group]


def top_negative_drivers(shap_row: pd.Series, n: int = 3) -> str:
    vals = [
        (feature.replace("shap_", ""), float(value))
        for feature, value in shap_row.items()
        if feature.startswith("shap_") and float(value) < 0
    ]
    vals = sorted(vals, key=lambda item: item[1])[:n]
    if not vals:
        return "none"
    return "; ".join(f"{k} ({v:.2f})" for k, v in vals)


def build_diagnosis() -> tuple[pd.DataFrame, pd.DataFrame]:
    panel = pd.read_csv(BASE / "data" / "processed" / "processed_panel.csv")
    forecast = pd.read_csv(TABLES / "table4_2025_target_gaps_and_risk_probabilities.csv")
    metrics = pd.read_csv(TABLES / "table3_rolling_origin_model_performance.csv")
    best_model = (
        metrics[(metrics["feature_set"] == "safe") & (metrics["horizon"] == 2)]
        .sort_values("mae")
        .iloc[0]["model"]
    )

    df, features = bep.make_supervised(panel, 2, "safe")
    train = df[df["target_year"] <= 2023].copy()
    pipe = bep.build_pipeline(bep.model_specs()[str(best_model)], features)
    pipe.fit(train[features], train["y_h2"])
    pre = pipe.named_steps["pre"]
    model = pipe.named_steps["model"]

    base = panel[panel["year"] == 2023].copy().sort_values("geo").reset_index(drop=True)
    x_trans = pre.transform(base[features])
    explainer = shap.TreeExplainer(model)
    shap_values = np.asarray(explainer.shap_values(x_trans))
    shap_df = pd.DataFrame(shap_values, columns=[f"shap_{feature}" for feature in features])
    shap_df.insert(0, "country", base["country"].to_numpy())
    shap_df.insert(0, "geo", base["geo"].to_numpy())
    shap_df.to_csv(TABLES / "table18_local_shap_2025_country_drivers.csv", index=False, encoding="utf-8-sig")

    diagnosis = forecast.merge(
        base[
            [
                "geo",
                "recycling_rate",
                "landfill_rate_proxy",
                "incineration_rate_proxy",
                "material_recycling_rate_proxy",
                "compost_digest_rate_proxy",
                "circular_material_use_rate",
                "resource_productivity",
                "recycling_rate_trend_3y",
                "landfill_rate_trend_3y",
            ]
        ],
        on="geo",
        how="left",
    ).merge(shap_df, on=["geo", "country"], how="left")

    groups = []
    group_scores = []
    top_drivers = []
    for _, row in diagnosis.iterrows():
        shap_row = row[[f"shap_{feature}" for feature in features if f"shap_{feature}" in row.index]]
        group, score = dominant_group(shap_row)
        groups.append(group)
        group_scores.append(score)
        top_drivers.append(top_negative_drivers(shap_row))

    diagnosis["severity_class"] = diagnosis.apply(severity, axis=1)
    diagnosis["dominant_bottleneck"] = groups
    diagnosis["dominant_bottleneck_score"] = group_scores
    diagnosis["top_local_risk_drivers"] = top_drivers
    diagnosis["diagnostic_profile"] = diagnosis["severity_class"] + " | " + diagnosis["dominant_bottleneck"]

    cols = [
        "geo",
        "country",
        "recycling_rate",
        "predicted_recycling_rate",
        "predicted_gap",
        "risk_probability",
        "ec_eea_2025_risk_label",
        "severity_class",
        "dominant_bottleneck",
        "top_local_risk_drivers",
        "landfill_rate_proxy",
        "incineration_rate_proxy",
        "material_recycling_rate_proxy",
        "compost_digest_rate_proxy",
        "circular_material_use_rate",
        "resource_productivity",
        "recycling_rate_trend_3y",
        "landfill_rate_trend_3y",
    ]
    diagnosis[cols].to_csv(TABLES / "table19_country_risk_diagnosis_profiles.csv", index=False, encoding="utf-8-sig")

    summary = (
        diagnosis.groupby(["severity_class", "dominant_bottleneck"], dropna=False)
        .agg(
            countries=("country", lambda x: ", ".join(x)),
            n=("country", "size"),
            mean_risk_probability=("risk_probability", "mean"),
            mean_gap=("predicted_gap", "mean"),
            mean_landfill_proxy=("landfill_rate_proxy", "mean"),
            mean_incineration_proxy=("incineration_rate_proxy", "mean"),
            mean_recycling_2023=("recycling_rate", "mean"),
        )
        .reset_index()
        .sort_values(["severity_class", "n"], ascending=[True, False])
    )
    summary.to_csv(TABLES / "table20_risk_group_bottleneck_summary.csv", index=False, encoding="utf-8-sig")

    high = diagnosis[diagnosis["risk_probability"] >= 0.5].copy()
    counts = high.groupby("dominant_bottleneck")["country"].count().sort_values(ascending=True)
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.barh(counts.index, counts.values, color="#1f4e79")
    ax.set_xlabel("Number of countries with risk probability >= 0.50")
    ax.set_title("Dominant local bottleneck among model-flagged high-risk countries")
    for i, v in enumerate(counts.values):
        ax.text(v + 0.05, i, str(v), va="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(FIGURES / "figure9_country_bottleneck_diagnosis.png", dpi=220)
    plt.close(fig)

    return diagnosis[cols], summary


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_table(doc: Document, caption: str, frame: pd.DataFrame, columns: list[str], max_rows: int | None = None) -> None:
    doc.add_paragraph(caption).runs[0].bold = True
    view = frame[columns].copy()
    if max_rows:
        view = view.head(max_rows)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = fmt(row[col], 3)
    doc.add_paragraph()


def build_doc(diagnosis: pd.DataFrame, summary: pd.DataFrame) -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("新版实验与贡献：EC-EEA 外部验证的 MSW Recycling Compliance Early-Warning and Treatment-Pathway Risk Diagnosis Framework")
    run.bold = True
    run.font.size = Pt(15)

    add_para(
        doc,
        "本次补强的核心变化：论文不再以“ML 预测 recycling rate”为主线，而是以“经过 EC-EEA 外部验证的 municipal solid waste recycling compliance early-warning and treatment-pathway risk diagnosis framework”为主线。新增实验是国家级 local SHAP 风险诊断、dominant bottleneck 分类和风险组汇总。",
    )

    doc.add_heading("一、新实验设计是否足够支撑新定位", level=1)
    design_rows = pd.DataFrame(
        [
            ["Compliance early warning", "EC-EEA external validation; strict information-cutoff validation", "证明框架能提前对齐官方风险标签，而不是只做内部回归预测。"],
            ["Target-gap decision output", "2025 target gap/risk probability; 2030/2035 stress projections", "证明框架能输出国家级 gap、risk 和优先级。"],
            ["Treatment-pathway risk diagnosis", "新增 table18 local SHAP; table19 country profiles; table20 bottleneck summary; figure9", "证明框架能指出每个国家的主导 bottleneck，而不是只给一个黑箱分数。"],
            ["Decision utility over simple baseline", "incremental value vs latest-rate baseline; Brier/F1/precision comparison", "证明 ML 框架的价值主要是校准和优先级排序，不夸大算法优势。"],
            ["Policy sensitivity and inertia", "scenario summary and country-level scenario table", "证明短期政策情景改善有限，支持 structural inertia 解释。"],
        ],
        columns=["模块", "实验", "支撑作用"],
    )
    add_table(doc, "表 1. 新定位下的实验设计重排", design_rows, list(design_rows.columns))

    doc.add_heading("二、新增国家级诊断实验结果", level=1)
    add_para(
        doc,
        "新增 table19_country_risk_diagnosis_profiles.csv 是关键补强。它把每个国家的 2025 risk probability、target gap、2023 treatment-pathway 指标、dominant local bottleneck 和 top local SHAP risk drivers 放在同一张表里。这样贡献可以写成 risk diagnosis framework，而不是普通预测模型。",
    )
    add_table(
        doc,
        "表 2. 高风险国家诊断样例，按 risk probability 排序。",
        diagnosis.sort_values("risk_probability", ascending=False),
        ["country", "predicted_gap", "risk_probability", "severity_class", "dominant_bottleneck", "top_local_risk_drivers", "landfill_rate_proxy", "incineration_rate_proxy"],
        max_rows=15,
    )
    add_table(
        doc,
        "表 3. 风险组与主导 bottleneck 汇总。",
        summary,
        ["severity_class", "dominant_bottleneck", "n", "mean_risk_probability", "mean_gap", "countries"],
        max_rows=20,
    )
    if (FIGURES / "figure9_country_bottleneck_diagnosis.png").exists():
        doc.add_picture(str(FIGURES / "figure9_country_bottleneck_diagnosis.png"), width=Inches(6.7))
        cap = doc.add_paragraph("Figure 9. Dominant local bottleneck among model-flagged high-risk countries.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("三、最终保留的四条强支撑贡献", level=1)
    contribution_rows = pd.DataFrame(
        [
            [
                "贡献 1：外部验证 early-warning 框架",
                "Table 5, Figure 4, Figure 8, Table 9",
                "EC-EEA AUC=0.981；0.50 threshold recall=1.000；2018 cutoff AUC=0.963。",
                "很强，应该作为第一贡献。",
            ],
            [
                "贡献 2：target-gap/risk probability 决策输出",
                "Table 4, Table 16, Table 17",
                "给出 EU-27 国家级 predicted rate、gap、risk probability 和优先级排序。",
                "强，符合 Waste Management 的 planning/decision-support 口味。",
            ],
            [
                "贡献 3：country-specific risk diagnosis beyond a single risk score",
                "新增 Table 18, Table 19, Table 20, Figure 9, SHAP Figure 5",
                "local SHAP + waste-system indicators 将国家风险区分为 recycling trajectory deficit、landfill lock-in、incineration/pathway pressure、circular-capacity/resource constraint 或 socioeconomic context。",
                "强支撑，但必须写成 country-specific diagnosis；不能写成 treatment pathway 是唯一或主导机制。",
            ],
            [
                "贡献 4：严格信息截断验证证明提前预警而非事后拟合",
                "Table 9, Figure 8, long-horizon Table 14",
                "2018-2023 cutoffs 均保持高 AUC/recall；7 年 horizon backtest 仍有 pseudo-target AUC。",
                "强，建议从 robustness 提升为主结果。",
            ],
        ],
        columns=["贡献", "强支撑图表", "关键证据", "期刊强度判断"],
    )
    add_table(doc, "表 4. 新定位下的强贡献与证据链", contribution_rows, list(contribution_rows.columns))

    doc.add_heading("四、论文中建议替换成的四条贡献表述", level=1)
    wordings = [
        "First, we develop an externally validated early-warning framework that translates public EU municipal solid waste indicators into compliance-risk probabilities benchmarked against the EC-EEA 2025 early-warning assessment.",
        "Second, we convert recycling forecasts into country-level target gaps, risk probabilities and priority rankings, allowing the framework to function as a decision-support tool rather than a generic recycling-rate predictor.",
        "Third, we add a country-specific diagnostic layer beyond a single risk score by combining local SHAP outputs with waste-system indicators, distinguishing whether predicted risk is mainly associated with recycling-trajectory deficits, landfill lock-in, incineration/pathway pressure, circular-capacity constraints, or socioeconomic context.",
        "Fourth, we use strict information-cutoff validation to show that the framework can recover official early-warning risk patterns using only pre-assessment information, reducing the concern that the model is merely reproducing retrospective labels.",
    ]
    for wording in wordings:
        add_para(doc, wording)

    doc.add_heading("五、必须避免的过度声称", level=1)
    for item in [
        "不要再把主贡献写成 ML predicts recycling rate。",
        "不要说 LightGBM/XGBoost 全面优于 baseline；应说在 external-label risk prioritisation 和 calibration 上有增量价值。",
        "不要把 SHAP/local SHAP 写成因果机制；只能写成 treatment-pathway risk diagnosis 或 fitted-model driver profile。",
        "不要把政策情景写成因果政策效果；只能写成 model-based sensitivity / structural inertia evidence。",
        "2030/2035 仍应写作 stress projection，不应写作已经验证的 long-term compliance forecast。",
    ]:
        add_para(doc, item)

    doc.save(OUT)
    print(OUT)


def main() -> None:
    diagnosis, summary = build_diagnosis()
    build_doc(diagnosis, summary)


if __name__ == "__main__":
    main()
