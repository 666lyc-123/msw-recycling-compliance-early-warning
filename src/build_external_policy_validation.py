from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE = Path(r"D:\VC code\MDPI2")
OUT_DIR = BASE / "outputs"
TABLE_CSV = OUT_DIR / "tables" / "table20_external_policy_report_validation.csv"
DOCX = OUT_DIR / "WasteManagement_external_policy_validation_planB.docx"


ROWS = [
    {
        "country": "Romania",
        "model_diagnosis": "High-gap high-risk; recycling-trajectory deficit",
        "official_evidence": "EEA reports no progress in municipal waste recycling during 2010-2022, a 2022 recycling rate of 12%, heavy reliance on landfilling, and a need to extend separate collection including bio-waste.",
        "source": "EEA 2025 Romania municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "Long-horizon recycling-system transition and landfill diversion remain central planning needs.",
    },
    {
        "country": "Malta",
        "model_diagnosis": "High-gap high-risk; recycling-trajectory deficit",
        "official_evidence": "EEA reports that Malta is at risk of missing 2025 municipal and packaging recycling targets and the 2035 landfill target, with no progress in municipal recycling or reducing landfilling since the previous early-warning report.",
        "source": "EEA 2025 Malta municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The model diagnosis is aligned with persistent recycling stagnation and landfill dependence.",
    },
    {
        "country": "Greece",
        "model_diagnosis": "High-gap high-risk; recycling-trajectory deficit",
        "official_evidence": "EEA reports that Greece is at risk of missing the 2025 municipal recycling target and the 2035 landfill target, with no progress since 2019 and landfilling stagnating at around 80%.",
        "source": "EEA 2025 Greece municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The country-level warning is supported by official evidence of stagnant recycling and persistent landfill dominance.",
    },
    {
        "country": "Cyprus",
        "model_diagnosis": "High-gap high-risk; recycling-trajectory deficit",
        "official_evidence": "EEA reports that Cyprus must accelerate progress toward 2025 municipal recycling targets, that municipal recycling has stagnated, and that only 15% of municipal waste was prepared for reuse or recycled in 2022.",
        "source": "EEA 2025 Cyprus municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The predicted high-risk profile is consistent with official evidence of low and stagnant recycling performance.",
    },
    {
        "country": "Portugal",
        "model_diagnosis": "High-gap high-risk; recycling-trajectory deficit",
        "official_evidence": "EEA reports that Portugal is at risk of missing 2025 municipal recycling targets and the 2035 landfill target, with progress on recycling and landfill reduction stagnating since 2016.",
        "source": "EEA 2025 Portugal municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The warning is consistent with official evidence that the target gap reflects structural stagnation rather than short-term noise.",
    },
    {
        "country": "Bulgaria",
        "model_diagnosis": "High-gap high-risk; recycling-trajectory deficit",
        "official_evidence": "EEA reports that Bulgaria has to speed up progress toward the 2025 recycling targets, has not improved municipal recycling since 2019, and needs better separate collection, especially for bio-waste.",
        "source": "EEA 2025 Bulgaria municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The model-based deficit diagnosis is supported by independent evidence on recycling stagnation and collection-system needs.",
    },
    {
        "country": "Hungary",
        "model_diagnosis": "High-gap high-risk; circular-capacity/resource constraint",
        "official_evidence": "EEA reports that Hungary is at risk of missing 2025 recycling targets and stresses the need to improve separate collection systems, increase awareness, use economic instruments, and further develop treatment infrastructure.",
        "source": "EEA 2025 Hungary municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The circular-capacity diagnosis is externally plausible because official evidence emphasizes collection and treatment infrastructure constraints.",
    },
    {
        "country": "Finland",
        "model_diagnosis": "High-risk near-target; incineration/pathway pressure",
        "official_evidence": "EEA reports that Finland is at risk of missing the 2025 municipal recycling target; landfill diversion has been accompanied by a significant increase in incineration while recycling has increased less.",
        "source": "EEA 2025 Finland municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The pathway-pressure diagnosis is supported by official evidence that energy recovery has absorbed landfill diversion more than recycling.",
    },
    {
        "country": "Austria",
        "model_diagnosis": "Stable achiever",
        "official_evidence": "EEA reports that Austria is on track for the 2025 municipal recycling target and reports a 2022 municipal recycling rate of 63%, above the 55% target, with landfill reduced to about 2%.",
        "source": "EEA 2025 Austria municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The low-risk model profile is consistent with official evidence of high recycling and very low landfill reliance.",
    },
    {
        "country": "Germany",
        "model_diagnosis": "Stable achiever",
        "official_evidence": "EEA reports that Germany's municipal recycling performance is above the 2025 target and that only about 1% of municipal waste is landfilled.",
        "source": "EEA 2025 Germany municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The low-risk classification is supported by official evidence of target-level recycling and minimal landfill reliance.",
    },
    {
        "country": "Netherlands",
        "model_diagnosis": "Stable achiever",
        "official_evidence": "EEA reports that the Netherlands is on track to meet 2025 municipal and packaging recycling targets as well as the 2035 landfill target, with progress in municipal recycling and reduced reliance on incineration.",
        "source": "EEA 2025 Netherlands municipal waste country profile",
        "consistency": "Strongly consistent",
        "planning_implication": "The low-risk profile is externally consistent with official evidence of target alignment and improving pathway balance.",
    },
]


def write_csv() -> None:
    TABLE_CSV.parent.mkdir(parents=True, exist_ok=True)
    with TABLE_CSV.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(ROWS[0].keys()))
        writer.writeheader()
        writer.writerows(ROWS)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def add_table(doc: Document) -> None:
    headers = ["Country", "Model diagnosis", "Independent official evidence", "Consistency", "Planning implication"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in ROWS:
        cells = table.add_row().cells
        values = [
            row["country"],
            row["model_diagnosis"],
            row["official_evidence"] + " Source: " + row["source"] + ".",
            row["consistency"],
            row["planning_implication"],
        ]
        for i, value in enumerate(values):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)


def write_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("External Policy-Report Validation of Country Diagnostic Profiles")
    run.bold = True
    run.font.size = Pt(14)

    add_para(
        doc,
        "This table provides an external plausibility check of selected country-specific diagnostic profiles using independent EEA municipal waste country profiles. It is not interpreted as causal validation; it tests whether the model-based diagnostic profiles are consistent with official policy and waste-system assessments.",
    )
    add_table(doc)
    doc.save(DOCX)


def main() -> None:
    write_csv()
    write_docx()
    print(TABLE_CSV)
    print(DOCX)


if __name__ == "__main__":
    main()
